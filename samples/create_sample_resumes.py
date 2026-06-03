"""
Generate sample DOCX resumes for testing the AI Resume Analyzer.
Run once after installing dependencies: python samples/create_sample_resumes.py
"""

from pathlib import Path

from docx import Document

OUTPUT_DIR = Path(__file__).parent / "resumes"
OUTPUT_DIR.mkdir(exist_ok=True)


def create_resume(filename: str, content: dict) -> None:
    """Create a DOCX resume from structured content."""
    doc = Document()
    doc.add_heading(content["name"], level=0)
    doc.add_paragraph(content["email"])
    doc.add_paragraph(content.get("phone", ""))

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(content["summary"])

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(content["skills"]))

    doc.add_heading("Experience", level=1)
    for exp in content["experience"]:
        doc.add_heading(exp["title"], level=2)
        doc.add_paragraph(f"{exp['company']} | {exp['duration']}")
        for bullet in exp["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Education", level=1)
    for edu in content["education"]:
        doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['year']})")

    doc.add_heading("Projects", level=1)
    for proj in content["projects"]:
        doc.add_heading(proj["name"], level=2)
        doc.add_paragraph(proj["description"])

    doc.save(OUTPUT_DIR / filename)
    print(f"Created: {OUTPUT_DIR / filename}")


# Sample Resume 1: Strong AI Engineer candidate
create_resume(
    "alice_chen_ai_engineer.docx",
    {
        "name": "Alice Chen",
        "email": "alice.chen@email.com",
        "phone": "+1 (555) 123-4567",
        "summary": (
            "AI Engineer with 4 years of experience building NLP and deep learning "
            "solutions. Expert in Python, PyTorch, and transformer models."
        ),
        "skills": [
            "Python", "PyTorch", "TensorFlow", "NLP", "Transformers",
            "Hugging Face", "LangChain", "RAG", "FastAPI", "Docker",
            "AWS", "Pandas", "NumPy", "scikit-learn", "Git", "SQL",
            "MLOps", "MLflow", "Prompt Engineering", "Fine-tuning",
        ],
        "experience": [
            {
                "title": "Senior AI Engineer",
                "company": "DataMind AI",
                "duration": "2022 - Present",
                "bullets": [
                    "Built RAG pipeline using LangChain and FAISS, improving answer accuracy by 35%",
                    "Fine-tuned BERT and GPT models for domain-specific NLP tasks",
                    "Deployed ML models via FastAPI on AWS with Docker containers",
                    "Implemented MLOps pipeline with MLflow for model versioning and monitoring",
                ],
            },
            {
                "title": "ML Engineer",
                "company": "InnovateTech",
                "duration": "2020 - 2022",
                "bullets": [
                    "Developed sentiment analysis model using PyTorch achieving 92% F1 score",
                    "Built data preprocessing pipelines with Pandas and Spark",
                    "Collaborated in Agile team to deliver AI features for SaaS product",
                ],
            },
        ],
        "education": [
            {
                "degree": "M.S. Computer Science (AI Specialization)",
                "institution": "Stanford University",
                "year": "2020",
            },
            {
                "degree": "B.Tech Computer Science",
                "institution": "IIT Delhi",
                "year": "2018",
            },
        ],
        "projects": [
            {
                "name": "LLM-Powered Document Q&A System",
                "description": (
                    "Built a RAG-based Q&A system using GPT-4, LangChain, and ChromaDB "
                    "for enterprise document search. Deployed with FastAPI and Docker."
                ),
            },
            {
                "name": "Real-time Sentiment Analysis API",
                "description": (
                    "Created a REST API using FastAPI and a fine-tuned BERT model "
                    "for real-time social media sentiment analysis."
                ),
            },
        ],
    },
)

# Sample Resume 2: Moderate match - Software developer transitioning to AI
create_resume(
    "bob_martinez_developer.docx",
    {
        "name": "Bob Martinez",
        "email": "bob.martinez@email.com",
        "phone": "+1 (555) 987-6543",
        "summary": (
            "Full-stack developer with 3 years of experience in Python and web development. "
            "Recently completed ML courses and built personal AI projects."
        ),
        "skills": [
            "Python", "JavaScript", "Flask", "Django", "SQL", "Git",
            "Pandas", "NumPy", "scikit-learn", "Machine Learning",
            "REST API", "Linux", "Docker", "MongoDB", "PostgreSQL",
        ],
        "experience": [
            {
                "title": "Python Developer",
                "company": "WebSolutions Inc.",
                "duration": "2021 - Present",
                "bullets": [
                    "Built REST APIs using Flask and Django for e-commerce platform",
                    "Implemented data analytics dashboard using Pandas and Matplotlib",
                    "Managed PostgreSQL databases and wrote complex SQL queries",
                    "Participated in Agile sprints and code reviews",
                ],
            },
            {
                "title": "Junior Developer Intern",
                "company": "StartupHub",
                "duration": "2020 - 2021",
                "bullets": [
                    "Developed backend services in Python",
                    "Created automated testing pipelines with CI/CD",
                ],
            },
        ],
        "education": [
            {
                "degree": "B.S. Information Technology",
                "institution": "State University",
                "year": "2020",
            },
        ],
        "projects": [
            {
                "name": "Movie Recommendation System",
                "description": (
                    "Built a collaborative filtering recommendation system using "
                    "scikit-learn and Pandas on the MovieLens dataset."
                ),
            },
            {
                "name": "Chatbot with Basic NLP",
                "description": (
                    "Created a rule-based chatbot with intent classification "
                    "using scikit-learn TF-IDF vectorizer."
                ),
            },
        ],
    },
)

# Sample Resume 3: Weak match - Data analyst with minimal AI experience
create_resume(
    "carol_williams_analyst.docx",
    {
        "name": "Carol Williams",
        "email": "carol.williams@email.com",
        "phone": "+1 (555) 456-7890",
        "summary": (
            "Data Analyst with 2 years of experience in business intelligence, "
            "reporting, and Excel-based analytics."
        ),
        "skills": [
            "Excel", "SQL", "Tableau", "Power BI", "Python", "Pandas",
            "Statistics", "Data Visualization", "Reporting", "Agile",
        ],
        "experience": [
            {
                "title": "Data Analyst",
                "company": "FinanceCorp",
                "duration": "2022 - Present",
                "bullets": [
                    "Created dashboards and reports using Tableau and Power BI",
                    "Wrote SQL queries for data extraction and analysis",
                    "Performed statistical analysis on financial datasets",
                    "Automated weekly reports using Python and Pandas",
                ],
            },
        ],
        "education": [
            {
                "degree": "B.A. Economics",
                "institution": "City College",
                "year": "2022",
            },
        ],
        "projects": [
            {
                "name": "Sales Dashboard",
                "description": (
                    "Built an interactive sales dashboard in Tableau "
                    "connecting to a PostgreSQL database."
                ),
            },
        ],
    },
)

print("\nAll sample resumes created successfully!")
print(f"Location: {OUTPUT_DIR.resolve()}")
